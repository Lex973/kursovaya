"""
Генератор защитного слова — 1 лист А4, шрифт 9pt.
"""

from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

ORIG = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Kursovaya.docx")
OUT  = os.path.join(os.path.dirname(os.path.abspath(__file__)), "ZashchitnoeSlovo.docx")

doc = Document(ORIG)

# Очищаем тело
for el in list(doc.element.body):
    if el.tag.split('}')[-1] != 'sectPr':
        doc.element.body.remove(el)

# Устанавливаем поля 1.5 cm со всех сторон (максимально компактно)
sec = doc.sections[0]
sec.left_margin   = Cm(1.5)
sec.right_margin  = Cm(1.5)
sec.top_margin    = Cm(1.2)
sec.bottom_margin = Cm(1.2)

_styles_map = {s.style_id: s for s in doc.styles}
S_NORMAL = _styles_map['Normal']

FONT_SIZE  = 9    # основной текст
TITLE_SIZE = 10   # заголовки блоков

def _p(text='', bold=False, size=FONT_SIZE, align=None,
       space_before=2, space_after=1, first_indent=False, italic=False):
    para = doc.add_paragraph()
    para.style = S_NORMAL
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.first_line_indent = Emu(450215) if first_indent else Pt(0)
    if align:
        para.alignment = align
    if text:
        r = para.add_run(text)
        r.bold   = bold
        r.italic = italic
        r.font.size = Pt(size)
    return para

def _h(text, size=TITLE_SIZE):
    """Жирный заголовок блока с линией-разделителем."""
    para = _p(text, bold=True, size=size, space_before=5, space_after=1)
    # Нижняя граница параграфа как разделитель
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4472C4')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para

def _li(text, size=FONT_SIZE):
    para = doc.add_paragraph()
    para.style = S_NORMAL
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after  = Pt(1)
    pf.left_indent  = Cm(0.4)
    pf.first_line_indent = Cm(-0.4)
    r = para.add_run('• ' + text)
    r.font.size = Pt(size)
    return para

def _table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))

    # Заголовок
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]
        c.text = h
        for pp in c.paragraphs:
            pp.paragraph_format.space_before = Pt(1)
            pp.paragraph_format.space_after  = Pt(1)
            pp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for rr in pp.runs:
                rr.bold = True
                rr.font.size = Pt(FONT_SIZE - 1)
        # Цвет заголовка
        tcPr = c._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'D9E2F3')
        shd.set(qn('w:val'),   'clear')
        tcPr.append(shd)

    # Данные
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri+1]
        for ci, val in enumerate(row_data):
            c = row.cells[ci]
            c.text = str(val)
            for pp in c.paragraphs:
                pp.paragraph_format.space_before = Pt(1)
                pp.paragraph_format.space_after  = Pt(1)
                pp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for rr in pp.runs:
                    rr.font.size = Pt(FONT_SIZE - 1)

    # Границы
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top','left','bottom','right'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'),   'single')
                b.set(qn('w:sz'),    '4')
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), 'AAAAAA')
                tcBorders.append(b)
            tcPr.append(tcBorders)

    # Ширина колонок
    if col_widths:
        for row in tbl.rows:
            for ci, w in enumerate(col_widths):
                if ci < len(row.cells):
                    row.cells[ci].width = Cm(w)

    _p(space_before=2, space_after=0)
    return tbl

# ═══════════════════════════════════════════════════════════════════
# ШАПКА
# ═══════════════════════════════════════════════════════════════════

_p('ЗАЩИТНОЕ СЛОВО', bold=True, size=12,
   align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=1)
_p('Куринин Л.В. | ИС-281 | Специальность 09.02.07',
   size=8, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=3)
_p('Тема: «Проектирование и разработка интернет-магазина косметики и парфюмерии»',
   bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)

# ═══════════════════════════════════════════════════════════════════
# 2 КОЛОНКИ через таблицу без границ
# ═══════════════════════════════════════════════════════════════════

layout = doc.add_table(rows=1, cols=2)
layout.style = _styles_map.get('a1', None) or layout.style  # убираем стиль

L = layout.rows[0].cells[0]   # левая колонка
R = layout.rows[0].cells[1]   # правая колонка
L.width = Cm(8.5)
R.width = Cm(8.5)

# ── Левая колонка ─────────────────────────────────────────────────

def lp(text='', bold=False, size=FONT_SIZE, space_before=2, space_after=1,
        first_indent=False, align=None, italic=False, cell=L):
    para = cell.add_paragraph()
    para.style = S_NORMAL
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.first_line_indent = Emu(450215) if first_indent else Pt(0)
    if align:
        para.alignment = align
    if text:
        r = para.add_run(text)
        r.bold   = bold
        r.italic = italic
        r.font.size = Pt(size)
    return para

def lh(text, cell=L):
    para = lp(text, bold=True, size=FONT_SIZE+1, space_before=5, space_after=1, cell=cell)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4472C4')
    pBdr.append(bottom)
    pPr.append(pBdr)

def lli(text, cell=L):
    para = cell.add_paragraph()
    para.style = S_NORMAL
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after  = Pt(1)
    pf.left_indent  = Cm(0.35)
    pf.first_line_indent = Cm(-0.35)
    r = para.add_run('• ' + text)
    r.font.size = Pt(FONT_SIZE)

# ─── Левая: О САЙТЕ ──────────────────────────────────────────────

lh('О сайте', L)
lp('Интернет-магазин косметики и парфюмерии — полнофункциональное '
   'SPA-приложение (React 19 + Laravel 13). Позволяет покупателям '
   'выбирать и заказывать товары онлайн, а администратору — управлять '
   'ассортиментом и заказами через отдельную панель.',
   first_indent=True, cell=L)

lh('Страницы сайта', L)
lli('Главная — hero-баннер, рекомендуемые товары (is_featured), бренды', L)
lli('Каталог — сетка товаров, боковые фильтры (категория, бренд, цена, пол, рейтинг, в наличии), поиск с debounce, пагинация', L)
lli('Карточка товара — галерея фото, атрибуты (из JSONB), рейтинг, кнопки «В корзину» / «В избранное», блок отзывов с формой', L)
lli('Корзина — список позиций, изменение количества (проверка остатка), удаление, итоговая сумма', L)
lli('Оформление заказа — форма доставки (ФИО, телефон, адрес), подтверждение', L)
lli('Личный кабинет — история заказов со статусами, избранное, редактирование профиля', L)
lli('Панель администратора — статистика (Recharts), CRUD товаров с загрузкой фото, управление заказами и пользователями', L)

lh('Путь покупателя', L)
lp('Регистрация → каталог + фильтры → карточка товара → корзина → '
   'оформление заказа → страница подтверждения с № заказа → '
   'история заказов в ЛК',
   first_indent=True, cell=L)

lh('База данных (10 таблиц)', L)
lli('users, categories, brands, products (JSONB attributes)', L)
lli('product_images, cart_items, orders, order_items, reviews, favorites', L)
lp('Цены фиксируются в order_items на момент оформления — снимок.',
   italic=True, cell=L, space_before=1, space_after=1)

lh('Тестирование', L)
lp('15 тест-кейсов, метод чёрного ящика — все пройдены.', cell=L, space_before=2)
lli('Среднее время API: 47 мс (норма ≤ 300 мс)', L)
lli('SQL-инъекции — нейтрализованы Eloquent ORM', L)
lli('Chrome 125 / Firefox 127 / Edge 125 — без ошибок', L)

# ── Правая колонка ────────────────────────────────────────────────

def rp(text='', **kw):
    kw.setdefault('cell', R)
    return lp(text, **kw)
def rh(text): lh(text, R)
def rli(text): lli(text, R)

rh('Технологический стек')

# Мини-таблица стека в правой колонке
tbl2 = R.add_table(rows=7, cols=2)
stack_rows = [
    ('Backend',   'Laravel 13 / PHP 8.3'),
    ('СУБД',      'PostgreSQL 16'),
    ('Auth',      'Laravel Sanctum (Bearer-токены)'),
    ('Frontend',  'React 19 + TypeScript 5'),
    ('Сборка',    'Vite 8'),
    ('CSS',       'Tailwind CSS 4'),
    ('State',     'Zustand v5 + TanStack Query v5'),
]
for ri, (k, v) in enumerate(stack_rows):
    c0, c1 = tbl2.rows[ri].cells[0], tbl2.rows[ri].cells[1]
    c0.width = Cm(1.8); c1.width = Cm(6.5)
    for c, txt, bld in [(c0, k, True), (c1, v, False)]:
        for pp in c.paragraphs:
            pp.clear()
            pp.paragraph_format.space_before = Pt(1)
            pp.paragraph_format.space_after  = Pt(1)
            pp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            rr = pp.add_run(txt)
            rr.bold = bld
            rr.font.size = Pt(FONT_SIZE - 0.5)
    # фон чётных строк
    if ri % 2 == 0:
        for c in (c0, c1):
            tcPr = c._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:fill'),  'EEF3FB')
            shd.set(qn('w:color'), 'auto')
            tcPr.append(shd)

rp(space_before=3, cell=R)

rh('Почему такой стек')
rli('Laravel — богатая экосистема (Sanctum, Eloquent), artisan-генераторы')
rli('PostgreSQL — JSONB для гибких атрибутов товаров, ACID-транзакции')
rli('React 19 + TypeScript — компонентный UI, статическая типизация')
rli('TanStack Query — кэш и авто-инвалидация запросов')
rli('Zustand persist — сессия сохраняется при перезагрузке')
rli('Vite 8 — мгновенный dev-сервер, нативные ES-модули')

rh('Ключевые маршруты API')
rp('Публичные:', bold=True, cell=R, space_before=3, space_after=0)
rli('GET /api/products — каталог с фильтрами и пагинацией')
rli('GET /api/products/{slug} — карточка товара')
rli('GET /api/categories, /api/brands')
rp('Покупатель (auth:sanctum):', bold=True, cell=R, space_before=2, space_after=0)
rli('POST /api/cart — добавить в корзину')
rli('POST /api/orders — оформить заказ')
rli('GET/POST/DELETE /api/favorites')
rli('POST /api/products/{id}/reviews')
rp('Администратор (+admin middleware):', bold=True, cell=R, space_before=2, space_after=0)
rli('GET /api/admin/stats — статистика')
rli('apiResource /api/admin/products — CRUD + фото')
rli('PATCH /api/admin/orders/{id} — смена статуса')
rli('GET/DELETE /api/admin/users')

rp(space_before=3, cell=R)
rp('Итого: 36 маршрутов | 10 таблиц БД | 15 тест-кейсов | 26 раб. дней',
   bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, cell=R)

# ═══════════════════════════════════════════════════════════════════
# Сохранение
# ═══════════════════════════════════════════════════════════════════

doc.save(OUT)
print(f'Готово: {OUT}')
